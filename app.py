import streamlit as st
import gspread
from google.oauth2.service_account import Credentials
from docx import Document
from io import BytesIO
import re
import json
import pandas as pd

# --- Konfigurasi Google Sheets ---
scope = ['https://spreadsheets.google.com/feeds', 'https://www.googleapis.com/auth/drive']
try:
    creds = Credentials.from_service_account_info(
        st.secrets["gcp_service_account"], scopes=scope
    )
    client = gspread.authorize(creds)
    gspread_auth_ok = True
except Exception as e:
    gspread_auth_ok = False
    st.error(f"Error saat otentikasi Google Sheets: {e}. Mode Google Sheets tidak akan berfungsi.")

# --- Fungsi untuk memproses dokumen Word ---
def process_docx(uploaded_file, data_map):
    try:
        doc_stream = BytesIO(uploaded_file.read())
        doc = Document(doc_stream)
        pattern = re.compile(r"\[(.*?:[A-Z]+[0-9]+)\]")
        
        def replace_in_container(container):
            full_text = ""
            for run in container.runs:
                full_text += run.text
            matches = pattern.findall(full_text)
            if not matches:
                return
            for match in matches:
                replacement_value = str(data_map.get(match, f"[{match}]"))
                full_text = full_text.replace(f"[{match}]", replacement_value)
            for run in container.runs:
                run.text = ""
            if container.runs:
                container.runs[0].text = full_text
            else:
                container.add_run(full_text)
                
        for p in doc.paragraphs:
            replace_in_container(p)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_in_container(p)
        
        new_doc_stream = BytesIO()
        doc.save(new_doc_stream)
        new_doc_stream.seek(0)
        return new_doc_stream
    except Exception as e:
        st.error(f"Error saat memproses dokumen: {e}")
        return None

# --- Fungsi Pembaca Excel ---
def read_excel_data(excel_file):
    try:
        xls = pd.ExcelFile(excel_file)
        sheet_names = xls.sheet_names
        all_data_dict = {}
        for sheet_name in sheet_names:
            df = pd.read_excel(excel_file, sheet_name=sheet_name, header=None)
            for r_idx, row in df.iterrows():
                for c_idx, cell_value in enumerate(row):
                    if pd.notna(cell_value):
                        col_letter = chr(ord('A') + c_idx)
                        cell_ref = f"{col_letter}{r_idx + 1}"
                        combined_key = f"{sheet_name}:{cell_ref}"
                        all_data_dict[combined_key] = cell_value
        return all_data_dict
    except Exception as e:
        st.error(f"Error saat membaca file Excel: {e}")
        return None

# --- Antarmuka Streamlit ---
st.title("Aplikasi Otomatisasi Dokumen Word Fleksibel 🚀")
st.markdown("Aplikasi ini akan mengganti tag dalam dokumen Word (.docx) dengan data dari Google Sheet atau file Excel yang Anda unggah.")

# Tambahkan pilihan mode
mode = st.radio("Pilih sumber data:", ("Google Sheets", "Unggah File Excel"))

tab1, tab2 = st.tabs(["Aplikasi", "Panduan Penggunaan"])

default_config_kompleks = """{
    "RINC SIMULATOR": "RINC SIMULATOR",
    "RINC CAPAIAN (TOTAL)": "RINC CAPAIAN (TOTAL)",
    "RINC JT IP": "RINC JT IP",
    "RINC SIMULATOR IP": "RINC SIMULATOR IP",
    "LAMP I-A (PROD JT)": "LAMP I-A (PROD JT)",
    "LAMP I-B ( JT IP)": "LAMP I-B ( JT IP)",
    "LAMP I-C (GUN.JT)": "LAMP I-C (GUN.JT)",
    "LAMP I-D (PROD FTD-FMS)": "LAMP I-D (PROD FTD-FMS)",
    "2-A Daftar Pers": "2-A Daftar Pers",
    "2-B STRUK JBTN": "2-B STRUK JBTN",
    "2-C DSP": "2-C DSP",
    "3-A Kesiapan Pesawat": "3-A Kesiapan Pesawat",
    "3-B HarPES": "3-B HarPES",
    "3-C Rinc Har": "3-C Rinc Har",
    "3-D BMP": "3-D BMP",
    "BMP Tahunan": "BMP Tahunan",
    "3-E Sucad": "3-E Sucad",
    "3-E Ranmor": "3-E Ranmor",
    "3-F Bliktek": "3-F Bliktek",
    "3-G Tools": "3-G Tools",
    "3-H GUN SUCAD": "3-H GUN SUCAD",
    "4 Alins-Alongins": "4 Alins-Alongins",
    "5 Inventaris Barang": "5 Inventaris Barang",
    "LAMP 6 (FASHARIN)": "LAMP 6 (FASHARIN)",
    "7-A LAMBANGJA": "7-A LAMBANGJA",
    "7-B Pot Hazard": "7-B Pot Hazard",
    "7-C Acc-Inc personel": "7-C Acc-Inc personel",
    "7-D Safety Meeting": "7-D Safety Meeting",
    "8 Permasalahan": "8 Permasalahan"
}"""
default_config_sederhana = """{
    "Data": "Data",
    "Obx": "Obx",
    "Control": "Control"
}"""

with tab1:
    if mode == "Google Sheets":
        st.header("Konfigurasi Data Google Sheet")
        if not gspread_auth_ok:
            st.warning("Mode ini tidak berfungsi karena terjadi kesalahan otentikasi.")
        st.markdown("Masukkan URL Google Sheet dan konfigurasi sheet dalam format JSON. Setiap entri adalah pasangan dari **'prefix tag'** dan **'nama sheet'**.")

        sheet_url_input = st.text_input(
            "URL Google Sheet", 
            "https://docs.google.com/spreadsheets/d/1d89txS35ZrBwk6_gyfOixvWZlvc149pgzhbekOka1uo/edit?usp=sharing"
        )
        sheet_config_choice = st.radio(
            "Pilih Konfigurasi",
            ("Sederhana", "Kompleks")
        )
        if sheet_config_choice == "Sederhana":
            sheet_config_input = st.text_area(
                "Konfigurasi Sheet (JSON)", 
                default_config_sederhana, 
                height=150
            )
        else:
            sheet_config_input = st.text_area(
                "Konfigurasi Sheet (JSON)", 
                default_config_kompleks, 
                height=400
            )

        uploaded_file_word = st.file_uploader("Pilih file Word (.docx) yang ingin diproses", type="docx")

        if st.button("Generate Dokumen"):
            if not uploaded_file_word:
                st.warning("Silakan unggah file .docx terlebih dahulu.")
            else:
                with st.spinner("Sedang memproses..."):
                    try:
                        sheet_config = json.loads(sheet_config_input)
                        all_data_dict = {}
                        for prefix, sheet_name in sheet_config.items():
                            try:
                                worksheet = client.open_by_url(sheet_url_input).worksheet(sheet_name)
                                all_values = worksheet.get_all_values()
                                for r_idx, row in enumerate(all_values):
                                    for c_idx, cell_value in enumerate(row):
                                        col_letter = chr(ord('A') + c_idx)
                                        cell_ref = f"{col_letter}{r_idx + 1}"
                                        combined_key = f"{prefix}:{cell_ref}"
                                        all_data_dict[combined_key] = cell_value
                            except gspread.WorksheetNotFound:
                                st.warning(f"Sheet dengan nama '{sheet_name}' tidak ditemukan. Melewati sheet ini.")
                                
                        if not all_data_dict:
                            st.error("Tidak ada data yang berhasil diambil dari Google Sheets. Mohon periksa URL dan nama sheet Anda.")
                        else:
                            processed_doc_stream = process_docx(uploaded_file_word, all_data_dict)
                            if processed_doc_stream:
                                file_name = uploaded_file_word.name.replace(".docx", "_generated.docx")
                                st.success("Dokumen berhasil dibuat!")
                                st.download_button(
                                    label="Unduh Dokumen Hasil",
                                    data=processed_doc_stream,
                                    file_name=file_name,
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                )
                    except json.JSONDecodeError:
                        st.error("Konfigurasi Sheet tidak valid. Mohon gunakan format JSON yang benar.")
                    except Exception as e:
                        st.error(f"Terjadi kesalahan: {e}")

    else: # Mode Unggah File Excel
        st.header("Unggah File Excel")
        st.info("Aplikasi akan membaca semua sheet dari file Excel yang diunggah. Gunakan nama sheet sebagai prefix tag di dokumen Word Anda.")

        uploaded_excel_file = st.file_uploader("Pilih file Excel (.xlsx)", type="xlsx")
        uploaded_word_file = st.file_uploader("Pilih file Word (.docx) yang ingin diproses", type="docx")

        if st.button("Generate Dokumen"):
            if not uploaded_excel_file or not uploaded_word_file:
                st.warning("Silakan unggah kedua file (Excel dan Word) terlebih dahulu.")
            else:
                with st.spinner("Sedang memproses..."):
                    data_dict = read_excel_data(uploaded_excel_file)
                    if data_dict:
                        processed_doc_stream = process_docx(uploaded_word_file, data_dict)
                        if processed_doc_stream:
                            file_name = uploaded_word_file.name.replace(".docx", "_generated.docx")
                            st.success("Dokumen berhasil dibuat!")
                            st.download_button(
                                label="Unduh Dokumen Hasil",
                                data=processed_doc_stream,
                                file_name=file_name,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )

with tab2:
    st.header("🚀 Panduan Singkat: Mengisi Dokumen Word Otomatis")
    st.markdown("Aplikasi ini membantu Anda mengisi data ke dalam dokumen Word secara otomatis dari dua sumber: **Google Sheets** atau **file Excel**.")

    st.subheader("Langkah 1: Siapkan Dokumen Word Anda")
    st.markdown("Buat template Word (`.docx`) dengan tag-tag data. Gunakan format **`[prefix:cell]`**.")
    st.markdown("- **`prefix`**: Nama yang Anda tentukan untuk mewakili sebuah sheet.")
    st.markdown("- **`cell`**: Nama sel yang ingin Anda ambil datanya (contoh: `A1`, `B2`).")
    st.markdown("**Contoh Tag**: `[Data:A1]`, `[Obx:A2]`, `[Control:A3]`")

    st.markdown("---")
    st.subheader("Langkah 2: Pilih Sumber Data dan Unggah File")

    st.markdown("**Mode Google Sheets**")
    st.markdown("1.  Pilih **'Google Sheets'** di bagian atas.")
    st.markdown("2.  Masukkan URL Google Sheet Anda.")
    st.markdown("3.  Pada bagian **'Konfigurasi Sheet'**, masukkan nama sheet yang ingin Anda gunakan. Prefix yang Anda masukkan di sini harus cocok dengan tag di dokumen Word Anda.")
    st.markdown("4.  Unggah file Word Anda.")

    st.markdown("**Mode Unggah File Excel**")
    st.markdown("1.  Pilih **'Unggah File Excel'** di bagian atas.")
    st.markdown("2.  Unggah file Excel (`.xlsx`) Anda.")
    st.markdown("    - **Catatan**: Di mode ini, **nama sheet** di file Excel Anda akan secara otomatis menjadi **prefix** tag. Jadi, tag seperti `[Data:A1]` akan mengambil data dari sheet bernama 'Data'.")
    st.markdown("3.  Unggah file Word Anda.")

    st.markdown("---")
    st.subheader("Langkah 3: Generate & Unduh Dokumen")
    st.markdown("1.  Klik tombol **'Generate Dokumen'**.")
    st.markdown("2.  Setelah proses selesai, tombol **'Unduh Dokumen Hasil'** akan muncul.")
    st.markdown("3.  Klik tombol tersebut untuk mengunduh dokumen yang sudah terisi data.")
